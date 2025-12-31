import io
import os
import random
import tempfile
from datetime import datetime
from typing import Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

from app.security import TokenData, require_role
from app.utils.logger import get_logger

# Setup logger for this module
logger = get_logger(__name__)

router = APIRouter(prefix="/api/documents", tags=["documents"])


def get_day_suffix(day: int) -> str:
    """
    Return the ordinal suffix for a given day (1st, 2nd, 3rd, 4th...).
    """
    if 11 <= day % 100 <= 13:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")


def format_date(date_str: str) -> str:
    """
    Format date string from YYYY-MM-DD to DD^th Mon YYYY format (e.g., 05th Nov 2025).
    """
    if not date_str:
        return ""

    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        day = date_obj.day
        suffix = get_day_suffix(day)
        return f"{day:02d}{suffix} {date_obj.strftime('%b %Y')}"
    except ValueError:
        logger.warning(f"Invalid date format: {date_str}, returning as-is")
        return date_str


def generate_reference_number(
    letter_date: Optional[str] = None,
    ref_month: Optional[str] = None,
    ref_year: Optional[str] = None,
) -> str:
    """
    Generate reference number in format 'Mon/YYYY-XYZ' where XYZ is a random 3-digit number.
    """
    date_obj = None

    # If user provided month and year, attempt to use them
    if ref_month and ref_year:
        try:
            month_int = int(ref_month)
            year_int = int(ref_year)
            date_obj = datetime(year=year_int, month=month_int, day=1)
        except ValueError:
            logger.warning(
                f"Invalid ref_month/ref_year for reference number: {ref_month}/{ref_year}, "
                "falling back to letter_date or current date"
            )

    if date_obj is None and letter_date:
        try:
            date_obj = datetime.strptime(letter_date, "%Y-%m-%d")
        except ValueError:
            logger.warning(f"Invalid letter_date for reference number: {letter_date}, using current date")

    if date_obj is None:
        date_obj = datetime.now()

    month_year = date_obj.strftime("%b/%Y")
    random_digits = random.randint(0, 999)
    return f"{month_year}-{random_digits:03d}"


def docx_to_pdf_reportlab(docx_path: str, replacements: dict) -> io.BytesIO:
    """
    Convert DOCX to PDF using ReportLab - cross-platform solution.
    """
    logger.info(f"Converting DOCX to PDF using ReportLab: {docx_path}")
    
    try:
        # Read the DOCX file
        doc = Document(docx_path)
        logger.info(f"DOCX file opened successfully: {len(doc.paragraphs)} paragraphs found")
    except Exception as e:
        logger.error(f"Failed to open DOCX file: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Could not read DOCX file: {str(e)}",
        ) from e

    # Create PDF in memory
    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18
    )

    # Container for the 'Flowable' objects
    elements = []

    # Define styles
    styles = getSampleStyleSheet()
    
    # Add custom styles with different alignments
    style_configs = [
        {
            'name': 'Left',
            'alignment': TA_LEFT,
            'fontSize': 11,
            'leading': 14,
            'spaceBefore': 6,
            'spaceAfter': 6
        },
        {
            'name': 'Center',
            'alignment': TA_CENTER,
            'fontSize': 12,
            'spaceAfter': 12
        },
        {
            'name': 'Right',
            'alignment': TA_RIGHT,
            'fontSize': 11,
            'leading': 14,
            'spaceBefore': 6,
            'spaceAfter': 6
        },
        {
            'name': 'Justify',
            'alignment': TA_JUSTIFY,
            'fontSize': 11,
            'leading': 14,
            'spaceBefore': 6,
            'spaceAfter': 6
        }
    ]
    
    for style_config in style_configs:
        try:
            styles.add(ParagraphStyle(
                name=style_config['name'],
                parent=styles['Normal'],
                alignment=style_config['alignment'],
                fontSize=style_config.get('fontSize', 11),
                leading=style_config.get('leading', 14),
                spaceBefore=style_config.get('spaceBefore', 0),
                spaceAfter=style_config.get('spaceAfter', 6)
            ))
        except (KeyError, ValueError):
            pass  # Style might already exist

    # Process paragraphs
    for paragraph in doc.paragraphs:
        # Get paragraph alignment BEFORE replacing placeholders
        # Try multiple ways to get alignment
        para_alignment = paragraph.alignment
        if para_alignment is None:
            # Try paragraph_format directly
            try:
                para_alignment = paragraph.paragraph_format.alignment
            except AttributeError:
                pass
        
        para_style_name = paragraph.style.name if paragraph.style else ""
        
        # Build text from runs to preserve any run-level formatting considerations
        text = paragraph.text

        # Replace placeholders - handle both {placeholder} and {{placeholder}} formats
        for key, value in replacements.items():
            text = text.replace(key, str(value))
            # Also handle double brace format
            if key.startswith('{') and not key.startswith('{{'):
                double_key = key.replace('{', '{{').replace('}', '}}')
                text = text.replace(double_key, str(value))

        if text.strip():
            # Determine alignment style based on DOCX alignment
            alignment_style = 'Left'  # Default to left
            
            # Try to get alignment from paragraph format XML if alignment is None
            if para_alignment is None:
                try:
                    # Check if paragraph format has alignment set in XML
                    p_format = paragraph.paragraph_format
                    if hasattr(p_format, '_element') and p_format._element is not None:
                        pPr = p_format._element.pPr
                        if pPr is not None:
                            jc = pPr.jc
                            if jc is not None and hasattr(jc, 'val') and jc.val is not None:
                                val = str(jc.val)
                                if 'right' in val.lower() or val == '2':
                                    para_alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    logger.debug(f"Found RIGHT alignment in XML: val='{val}'")
                                elif 'center' in val.lower() or 'centre' in val.lower() or val == '1':
                                    para_alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    logger.debug(f"Found CENTER alignment in XML: val='{val}'")
                                elif 'left' in val.lower() or val == '0':
                                    para_alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    logger.debug(f"Found LEFT alignment in XML: val='{val}'")
                                elif 'justify' in val.lower() or val == '3':
                                    para_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    logger.debug(f"Found JUSTIFY alignment in XML: val='{val}'")
                    
                    # Also check style's paragraph format alignment
                    if para_alignment is None and paragraph.style:
                        try:
                            style_alignment = paragraph.style.paragraph_format.alignment
                            if style_alignment is not None:
                                para_alignment = style_alignment
                                logger.debug(f"Found alignment from style: {style_alignment}")
                        except (AttributeError, Exception):
                            pass
                except (AttributeError, Exception) as e:
                    logger.debug(f"Could not read alignment from XML: {e}")
            
            # Check explicit paragraph alignment first
            if para_alignment == WD_ALIGN_PARAGRAPH.CENTER:
                alignment_style = 'Center'
            elif para_alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                alignment_style = 'Right'
            elif para_alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                alignment_style = 'Justify'
            elif para_alignment == WD_ALIGN_PARAGRAPH.LEFT:
                alignment_style = 'Left'
            elif para_alignment is None:
                # When alignment is None, it inherits from style or defaults to left
                # Check style name for alignment hints
                style_lower = para_style_name.lower()
                if 'right' in style_lower:
                    alignment_style = 'Right'
                    logger.debug(f"Using Right alignment from style name: {para_style_name}")
                elif 'center' in style_lower or 'centre' in style_lower:
                    alignment_style = 'Center'
                    logger.debug(f"Using Center alignment from style name: {para_style_name}")
                else:
                    alignment_style = 'Left'  # Default
            
            # Log alignment detection for debugging
            if alignment_style != 'Left' or para_alignment is not None:
                logger.info(f"Paragraph alignment: text='{text[:50]}...', para_alignment={para_alignment}, style_name='{para_style_name}', final='{alignment_style}'")
            
            # Detect if it's a heading for bold formatting
            is_heading = (
                para_style_name.startswith('Heading') or
                '[' in text or
                'EXPERIENCE' in text.upper() or
                'Sub:' in text or
                (text.strip().startswith('To') and len(text.strip()) < 5)
            )

            # Apply formatting and alignment
            if is_heading:
                formatted_text = f"<b>{text}</b>"
            else:
                formatted_text = text
            
            # Use the determined alignment style
            style_to_use = styles.get(alignment_style, styles['Left'])
            p = Paragraph(formatted_text, style_to_use)
            
            elements.append(p)
            elements.append(Spacer(1, 12))

    # Process tables if any
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            cell_alignments = []
            for cell in row.cells:
                # Process each paragraph in the cell
                cell_paragraphs = []
                for para in cell.paragraphs:
                    if para.text.strip():
                        # Get alignment for each paragraph
                        cell_alignment = para.alignment
                        # If None, try to infer from style
                        if cell_alignment is None:
                            para_style_name = para.style.name if para.style else ""
                            style_lower = para_style_name.lower()
                            if 'right' in style_lower:
                                cell_alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif 'center' in style_lower or 'centre' in style_lower:
                                cell_alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell_paragraphs.append((para.text.strip(), cell_alignment))
                
                if cell_paragraphs:
                    # Use first paragraph's alignment for the cell
                    cell_text = " ".join(text for text, _ in cell_paragraphs)
                    cell_alignment = cell_paragraphs[0][1] if cell_paragraphs else None
                    row_texts.append(cell_text)
                    cell_alignments.append(cell_alignment)
            
            if any(row_texts):
                # Process each cell separately to preserve alignment
                for i, cell_text in enumerate(row_texts):
                    if cell_text.strip():
                        # Replace placeholders in cell text
                        for key, value in replacements.items():
                            cell_text = cell_text.replace(key, str(value))
                            if key.startswith('{') and not key.startswith('{{'):
                                double_key = key.replace('{', '{{').replace('}', '}}')
                                cell_text = cell_text.replace(double_key, str(value))
                        
                        # Determine alignment style
                        cell_alignment = cell_alignments[i] if i < len(cell_alignments) else None
                        alignment_style = 'Left'  # Default
                        if cell_alignment == WD_ALIGN_PARAGRAPH.CENTER:
                            alignment_style = 'Center'
                        elif cell_alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                            alignment_style = 'Right'
                        elif cell_alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                            alignment_style = 'Justify'
                        
                        style_to_use = styles.get(alignment_style, styles['Left'])
                        p = Paragraph(cell_text, style_to_use)
                        elements.append(p)
                elements.append(Spacer(1, 6))

    # Build PDF
    try:
        pdf.build(elements)
        buffer.seek(0)
        logger.info(f"PDF generation completed successfully: {len(buffer.getvalue())} bytes")
        return buffer
    except Exception as e:
        logger.error(f"Error building PDF with ReportLab: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error converting DOCX to PDF: {str(e)}",
        ) from e


def build_replacements(
    name: str,
    role: str,
    start_date: str,
    end_date: str,
    letter_date: str = "",
    address: str = "",
    phone_number: str = "",
    ref_month: Optional[str] = None,
    ref_year: Optional[str] = None,
) -> Dict[str, str]:
    """
    Build replacement dictionary with properly formatted text.
    """
    logger.info(f"Building replacements for: name={name}, role={role}, start_date={start_date}, end_date={end_date}")
    
    # Format dates properly
    formatted_start_date = format_date(start_date)
    formatted_end_date = format_date(end_date)
    formatted_letter_date = (
        format_date(letter_date) if letter_date else format_date(datetime.now().strftime("%Y-%m-%d"))
    )
    reference_number = generate_reference_number(
        letter_date if letter_date else None,
        ref_month=ref_month,
        ref_year=ref_year,
    )
    
    # Clean and format name (proper case)
    clean_name = name.strip()
    if clean_name:
        clean_name = ' '.join(word.capitalize() for word in clean_name.split())
    
    # Clean and format role (proper case)
    clean_role = role.strip()
    if clean_role:
        clean_role = ' '.join(word.capitalize() for word in clean_role.split())
    
    # Build replacements - support both {placeholder} and {{placeholder}} formats
    replacements = {
        "{name}": clean_name,
        "{role}": clean_role,
        "{start_date}": formatted_start_date,
        "{end_date}": formatted_end_date,
        "{letter_date}": formatted_letter_date,
        "{address}": address.strip() if address else "",
        "{phone_number}": phone_number.strip() if phone_number else "N/A",
        # Double brace format
        "{{name}}": clean_name,
        "{{role}}": clean_role,
        "{{start_date}}": formatted_start_date,
        "{{end_date}}": formatted_end_date,
        "{{letter_date}}": formatted_letter_date,
        "{{address}}": address.strip() if address else "",
        "{{phone_number}}": phone_number.strip() if phone_number else "N/A",
        "{ref.no}": reference_number,
        "{{ref.no}}": reference_number,
    }
    
    logger.info(f"Built replacements: {list(replacements.keys())}")
    return replacements


@router.post("/experience-letter/generate")
async def generate_experience_letter(
    name: str = Form(...),
    role: str = Form(...),
    start_date: str = Form(...),
    end_date: str = Form(...),
    template_file: Optional[UploadFile] = File(None),
    address: str = Form(""),
    phone: str = Form(""),
    signatory_name: str = Form(""),
    letter_date: Optional[str] = Form(None),
    ref_month: Optional[str] = Form(None),
    ref_year: Optional[str] = Form(None),
    current_user: TokenData = Depends(require_role("Admin", "Super Admin")),
):
    """
    Generate experience letter PDF from DOCX template with placeholders replaced.
    If no template is provided, returns an error (template is required in this version).
    """
    logger.info(f"Experience letter generation requested by user {current_user.username} (ID: {current_user.user_id})")
    logger.info(f"Input parameters: name={name}, role={role}, start_date={start_date}, end_date={end_date}")
    
    del current_user  # role validation handled by dependency

    if not name.strip() or not role.strip():
        logger.warning("Experience letter generation failed: Missing required fields (name or role)")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Candidate name and role are required.",
        )

    # Template file is required in this version
    if template_file is None:
        logger.warning("Experience letter generation failed: No template file provided")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Template file (DOCX) is required. Please upload a DOCX template file.",
        )

    # Validate file type
    filename = template_file.filename or ""
    if not filename.lower().endswith('.docx'):
        logger.warning(f"Invalid file type: {filename}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only .docx files are supported. Please upload a DOCX template file.",
        )

    try:
        # Create temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            # Save uploaded file
            input_path = os.path.join(temp_dir, "template.docx")
            template_bytes = await template_file.read()

            if not template_bytes:
                logger.warning("Experience letter generation failed: Empty template file")
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Uploaded template is empty.",
                )

            logger.info(f"Template file read successfully: {len(template_bytes)} bytes")

            with open(input_path, "wb") as f:
                f.write(template_bytes)

            # Prepare replacements dictionary
            replacements = build_replacements(
                name=name,
                role=role,
                start_date=start_date,
                end_date=end_date,
                letter_date=letter_date if letter_date else "",
                address=address,
                phone_number=phone,
                ref_month=ref_month,
                ref_year=ref_year,
            )

            # Convert DOCX to PDF using ReportLab
            logger.info("Starting DOCX to PDF conversion with ReportLab")
            pdf_buffer = docx_to_pdf_reportlab(input_path, replacements)

            # Return the PDF file
            output_size = len(pdf_buffer.getvalue())
            logger.info(f"Experience letter generated successfully: {output_size} bytes")

            headers = {
                "Content-Disposition": f'attachment; filename="Experience_Letter_{name.replace(" ", "_")}.pdf"',
                "Cache-Control": "no-store",
            }
            return StreamingResponse(
                pdf_buffer,
                media_type="application/pdf",
                headers=headers
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating experience letter: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error generating experience letter: {str(e)}",
        ) from e
